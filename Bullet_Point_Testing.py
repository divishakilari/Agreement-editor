import streamlit as st
from docxtpl import DocxTemplate
import base64
import docx
from PIL import Image
import os
from docx.shared import Pt
import subprocess
from subprocess import Popen
from docx.shared import RGBColor
# import sys
st.set_page_config(layout="wide")
subprocess.call(["apt", "install", "libreoffice", "--no-install-recommends"])
# import os

path = os.getcwd()
path1 = path + "/Documents"
path2 = path + "/generated documents"
path3 = path + "/edited documents"

image = Image.open(path+'/Payatu_logo.webp')
st.sidebar.image(image, width=280)

if "count" not in st.session_state:
    st.session_state["count"] = 0
if "add_doc" not in st.session_state:
    st.session_state["add_doc"] = docx.Document()
if "add_docR" not in st.session_state:
    st.session_state["add_docR"] = docx.Document()
if "additional clause" not in st.session_state:
    st.session_state["additional clause"] = ""

def addBulletR(document, text, cursor_word):
    for para in range(len(document.paragraphs)):
        if document.paragraphs[para].text == cursor_word:
            cursor_paragraph = document.paragraphs[para]
            new_paragraph = cursor_paragraph.insert_paragraph_before('',style = "List Bullet 3").add_run(text)
            new_paragraph.font.size = Pt(12)
            new_paragraph.font.name = 'Times New Roman'
            new_paragraph.font.color.rgb = RGBColor(0, 112, 192)
            break
    document.save(path3+"/Edited Agreement.docx")
def addBullet(document, text, cursor_word):
    for para in range(len(document.paragraphs)):
        if document.paragraphs[para].text == cursor_word:
            cursor_paragraph = document.paragraphs[para]
            new_paragraph = cursor_paragraph.insert_paragraph_before('',style = "List Bullet 3").add_run(text)
            new_paragraph.font.size = Pt(12)
            new_paragraph.font.name = 'Times New Roman'
            break
    document.save(path3+"/Edited Agreement.docx")

def addClauseR(document, text, cursor_word):
    for para in range(len(document.paragraphs)):
        if document.paragraphs[para].text == cursor_word:
            cursor_paragraph = document.paragraphs[para]
            new_paragraph = cursor_paragraph.insert_paragraph_before('',style = "List 2").add_run(text)
            new_paragraph.font.size = Pt(12)
            new_paragraph.font.name = 'Times New Roman'
            new_paragraph.font.color.rgb = RGBColor(0, 112, 192)
            break
    document.save(path3+"/Edited AgreementR.docx")

def addClause(document, text, cursor_word):
    for para in range(len(document.paragraphs)):
        if document.paragraphs[para].text == cursor_word:
            cursor_paragraph = document.paragraphs[para]
            new_paragraph = cursor_paragraph.insert_paragraph_before('',style = "List 2").add_run(text)
            new_paragraph.font.size = Pt(12)
            new_paragraph.font.name = 'Times New Roman'
            break
    document.save(path3+"/Edited Agreement.docx")
    

def convert_to_pdf(input_docx, out_folder):
    p = Popen(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    
#     p = Popen(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
#                out_folder, input_docx])
    
    p.communicate()


def displayPDF(file):
    #Opening file from file path
    with open(file, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = F'<embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf">'
   # Displaying File
    st.markdown(pdf_display, unsafe_allow_html=True)


def triggerfunction(file, file_name):
    convert_to_pdf(file, path2+'/')
    with open(path2+"/"+file_name, "rb") as pdf_file:
        PDFbyte = pdf_file.read()

    st.download_button(label="Download PDF",
                       data=PDFbyte,
                       file_name= file_name+".pdf",
                       mime='application/octet-stream')

def GeneralServiceAgreement(path1, path2, path3):
    st.markdown(""" <style> .font {
        font-size:35px ; font-family: 'Consolas'; color: #0077b6; text-align: center} 
        </style> """, unsafe_allow_html=True)
    st.markdown('<p class="font">GENERAL SERVICE AGREEMENT</p>',
                unsafe_allow_html=True)
    #----------------from here--------------------------------------# 
    if(st.session_state["count"] == 0):
        st.session_state["add_docR"] = docx.Document(path1 + "/General Service AgreementR.docx")
        st.session_state["add_doc"] = docx.Document(path1 + "/General Service Agreement.docx")
    else:
        st.session_state["add_docR"] = docx.Document(path3 + "/Edited AgreementR.docx")
        st.session_state["add_doc"] = docx.Document(path3 + "/Edited Agreement.docx")
    st.session_state["additional areas"] = st.sidebar.text_input("Type/paste the extra areas in here without any serial number ex: Scheduling Team meeings,")
    if st.sidebar.button("Option to Add Additional areas"):
       if st.session_state["additional areas"] is not None:
           st.session_state["count"] = 1
           addBulletR(st.session_state["add_docR"], st.session_state["additional areas"],"{{ Add_Additional_Areas }}")
           addBullet(st.session_state["add_doc"], st.session_state["additional areas"],"{{ Add_Additional_Areas }}")
    
    st.session_state["additional clause"] = st.sidebar.text_input("Type/paste the extra clause in here along with serial number ex: 25. additional clause")
    if st.sidebar.button("Option to Add Additional Clauses"):
        if st.session_state["additional clause"] is not None:
            st.session_state["count"] = 1
            addClauseR(st.session_state["add_docR"], st.session_state["additional clause"], "{{ Option_to_add_more_clauses }}")
            addClause(st.session_state["add_doc"], st.session_state["additional clause"],"{{ Option_to_add_more_clauses }}")
    
    st.session_state["add_doc"].save(path3+"/Edited Agreement.docx")
    st.session_state["add_docR"].save(path3+"/Edited AgreementR.docx")
    if(st.session_state["count"] != 0):
        doc = DocxTemplate(path3+"/Edited AgreementR.docx")
        doc1 = DocxTemplate(path3+"/Edited Agreement.docx")
    else:
        doc = DocxTemplate(path1+"/General Service AgreementR.docx")
        doc1 = DocxTemplate(path1+"/General Service Agreement.docx")
#-------------To here-----------------------------------------------------#


    # doc = DocxTemplate(path1+"/General Service AgreementR.docx")
    # doc1 = DocxTemplate(path1+"/General Service Agreement.docx")
    context = {"Place": st.sidebar.text_input("Place:", value="Place"),
               "dd_mm_yy": st.sidebar.date_input("Date:"),
               "Client1": st.sidebar.text_input("Client Name:", value="Client"),
               "Address_cl": st.sidebar.text_input("Client Address:", value="Client Address"),
               "Sole_Proprietor_or_Partner_or_Duly_Authorized_Member_Of_Staff_or_NA_cl": st.sidebar.selectbox("Client Represented By",
                                                                                                              ("Sole Proprietor", "Partner", "Duly Authorized Member of Staff")),
               "Mr_or_Ms_cl": st.sidebar.selectbox("Title of Client Representor:", ("Mr.", "Mrs.")),
               "Client_Representative": st.sidebar.text_input("Client Representative", value="Client Representative"),
               "Contractor1": st.sidebar.text_input("Contractor Name:", value="Contractor Name"),
               "Address_co": st.sidebar.text_input("Contractor Address:", value="Contractor Address"),
               "Sole_Proprietor_or_Partner_or_Duly_Authorized_Member_Of_Staff_or_NA_co": st.sidebar.selectbox("Contractor Represented By",
                                                                                                              ("Sole Proprietor", "Partner", "Duly Authorized Member of Staff")),
               "Mr_or_Ms_co": st.sidebar.selectbox("Title of Contractor represntative:", ("Mr.", "Mrs.")),
               "Contractor_Representative": st.sidebar.text_input("", value="Contractor Representative"),
               "Goods": st.sidebar.text_input("Name of the Goods to be supplied", value="Name of the Goods to be supplied"),

               "Purpose": st.sidebar.text_input("Purpose", value="Purpose"),
               "From_date": st.sidebar.date_input("Effective from date"),

               "To_Date": st.sidebar.date_input("Effective till date"),
               "Duration": st.sidebar.text_input("Duration", value="Duration"),

               "Service_1": st.sidebar.text_input("Services provided by Contractor to Client", value="Services provided by Contractor to Client"),
               # "Add_any_Additional_Services": st.sidebar.text_input("Add if there are any additional services else leave it empty"),
               "A_flat_fee_or_In_installments_or_Other_Consideration": st.sidebar.selectbox("Payment Method", ("Down/Flat Payment", "Installments", "Other Considerations")),
               "AmountF": st.sidebar.text_input("In case of Flat Fee: (Enter Amount)", value="Flat Fee Amount"),
               "Amount_of_Installments": st.sidebar.text_input("Incase of Installments: (Enter Amount)", value="Installments Amount"),
               "AmountI": st.sidebar.text_input("First Installment Amount: ", value="Enter Amount"),
               "AmountS": st.sidebar.text_input("Second Installment Amount: ", value="Enter Amount"),
               "Additional_Installments": st.sidebar.text_input("Additional Installmetns(if any) or leave it empty: "),
               "Consideration": st.sidebar.text_input("Consideration: "),
               "Before_or_After_or_During_or_In_Installments": st.sidebar.selectbox("sevice by contractor must be given: ", ("Before", "After", "During / In Installments")),
               "method_of_payment": st.sidebar.selectbox("Method of Payment: ", ("Debit Card", "Credit Card", "Cash")),

               "Amount_cl": st.sidebar.text_input("Amount client need to pay", value="Amount Client need to pay"),
               "Client_or_Contractor_or_Both_Parties": st.sidebar.selectbox("Material developed /produced will be the property of: ", ("Client", "Contractor", "Both Parties")),
               "Amount_Re": st.sidebar.text_input("Client reimburse to extent of: ", value="Amount Reimbursed"),
               "Number_of_Days": st.sidebar.text_input("Claims must be settled within", value="N days"),
               "Name_of_State_or_District": st.sidebar.text_input("State/District", value="State/District"),
               "Number": st.sidebar.text_input("Duration N days: ", value="Duration (in numbers)"),
               "Court_of_Law_or_Arbitral_tribunal": st.sidebar.selectbox("May seek compensation in appropriate: ", ("Court of Law", "Arbitral Tribunal")),
               # "Additional Clauses": st.sidebar.text_input("Add Additional Clause if any / Leave it empty"),

               "dd_mm_yy_1": st.sidebar.date_input("Date of sigining the agreement"),


               "Client": st.sidebar.text_input("", value="Client Name", placeholder="Client Name"),
               "Client_Representative_Name": st.sidebar.text_input("Client Representative Name", value="Client Representative Name"),
               "Client_Representative_Position": st.sidebar.text_input("Client Representative Position", value="Client Representative Position"),


               "Contractor": st.sidebar.text_input("Contractor Name", value="Contractor Name"),
               "Contractor_Representative_Name": st.sidebar.text_input("Contractor Representative Name", value="Contractor Representative Name"),
               "Contractor_Representative_Position": st.sidebar.text_input("Contractor Representative Position", value="Contractor Representative Position"),
               "Contractor_Representative_Signature": st.sidebar.text_input("Contractor Representative Signature", value="Contractor Representative Signature"),
               "Witness_1_Name": st.sidebar.text_input("Witness Name", value="Witness Name"),


               "Witness_2_Name": st.sidebar.text_input("Witness Name 2:", value="Witness Name 2 or NA"),
               "Add_Additional_Areas": "{{ Add_Additional_Areas }}",
               "Option_to_add_more_clauses": "{{ Option_to_add_more_clauses }}"
               }
    doc.render(context)
    doc1.render(context)
    doc.save(path2+"/General Service agreement generatedR.docx")
    doc1.save(path2+"/General Service agreement generated.docx")
    convert_to_pdf(path2+"/General Service agreement generatedR.docx", path2+'/')
    f1 = path2+'/General Service agreement generatedR.pdf'

    if st.button('Done editing'):
        final_doc = DocxTemplate(path2+"/General Service agreement generated.docx")
        context1 = {}
        final_doc.render(context1)
        final_doc.save(path2+"/General Service agreement generated.docx")
        triggerfunction(path2+"/General Service agreement generated.docx","General Service agreement generated.pdf")
    displayPDF(f1)

    
def NullconGoaSponsorshipAgreement(path1, path2, path3):
    st.markdown(""" <style> .font {
        font-size:35px ; font-family: 'Consolas'; color: #0077b6; text-align: center} 
        </style> """, unsafe_allow_html=True)
    st.markdown('<p class="font">NULLCON Goa Sponsorship Agreement</p>',
                unsafe_allow_html=True)
    #----------------from here--------------------------------------# 
    if(st.session_state["count"] == 0):
        st.session_state["add_docR"] = docx.Document(path1 + "/Nullcon Goa Sponsorship AgreementR.docx")
        st.session_state["add_doc"] = docx.Document(path1 + "/Nullcon Goa Sponsorship Agreement.docx")
    else:
        st.session_state["add_docR"] = docx.Document(path3 + "/Edited AgreementR.docx")
        st.session_state["add_doc"] = docx.Document(path3 + "/Edited Agreement.docx")
   
    # st.session_state["additional areas"] = st.sidebar.text_input("Type/paste the extra areas in here without any serial number ex: Scheduling Team meeings,")
    # if st.sidebar.button("Option to Add Additional areas"):
    #    if st.session_state["additional areas"] is not None:
    #        st.session_state["count"] = 1
    #        addBulletR(st.session_state["add_docR"], st.session_state["additional areas"],"{{ Add_Additional_Areas }}")
    #        addBullet(st.session_state["add_doc"], st.session_state["additional areas"],"{{ Add_Additional_Areas }}")
        
    st.session_state["Add Benefit"] = st.sidebar.text_input("Type the additional benefits here and then click the ADD button")
    if st.sidebar.button("Option to Add Benefit"):
        if st.session_state["Add Benefit"] is not None:
            st.session_state["count"] = 1
            addBulletR(st.session_state["add_docR"], st.session_state["Add Benefit"],"{{ Add_another_benefit }}")
            addBullet(st.session_state["add_doc"], st.session_state["Add Benefit"],"{{ Add_another_benefit }}")
    
    
    st.session_state["add_doc"].save(path3+"/Edited Agreement.docx")
    st.session_state["add_docR"].save(path3+"/Edited AgreementR.docx")
    if(st.session_state["count"] != 0):
        doc = DocxTemplate(path3+"/Edited AgreementR.docx")
        doc1 = DocxTemplate(path3+"/Edited Agreement.docx")
    else:
        doc = DocxTemplate(path1+"/Nullcon Goa Sponsorship AgreementR.docx")
        doc1 = DocxTemplate(path1+"/Nullcon Goa Sponsorship Agreement.docx")
#-------------To here-----------------------------------------------------#


    # doc = DocxTemplate(path1+"/General Service AgreementR.docx")
    # doc1 = DocxTemplate(path1+"/General Service Agreement.docx")
    context = {"dd_mm_yy": st.sidebar.date_input("Date:"),
               "Sponsorship_Name": st.sidebar.text_input("Sponsorship Name:", value="Goodie Bag Sponsorship Booth"),
               "Sponsor_Name": st.sidebar.text_input("Sponsor Name:", value="Sponsor Name"),
               "Sponser_Address": st.sidebar.text_input("Sponsor Address:", value="Sponsor Address"),
               "Conference_Dates": st.sidebar.text_input("Conference_Dates", value="20 sep to 22 sep"),
               "Sponsorship_Level": st.sidebar.text_input("Sponsorship_Level", value="Goodie Bag Sponsorship + Standard Exhibition Booth"),
               "Sponsorship_Fees": st.sidebar.text_input("Sponsorship_Fees", value="$200 + GST"),
                "Benefit_1": st.sidebar.text_input("Benefit", value="Add Benefit 1"),
               "Add_another_benefit": "{{ Add_another_benefit }}",
               
               # "Add_Additional_Areas": "{{ Add_Additional_Areas }}"

               }
    doc.render(context)
    doc1.render(context)
    doc.save(path2+"/Nullcon Goa Sponsorship Agreement generatedR.docx")
    doc1.save(path2+"/Nullcon Goa Sponsorship Agreement generated.docx")
    convert_to_pdf(path2+"/Nullcon Goa Sponsorship Agreement generatedR.docx", path2+'/')
    f1 = path2+'/Nullcon Goa Sponsorship Agreement generatedR.pdf'

    if st.button('Done editing'):
        final_doc = DocxTemplate(path2+"/Nullcon Goa Sponsorship Agreement generated.docx")
        context1 = {}
        final_doc.render(context1)
        final_doc.save(path2+"/Nullcon Goa Sponsorship Agreement generated.docx")
        triggerfunction(path2+"/Nullcon Goa Sponsorship Agreement generated.docx","Nullcon Goa Sponsorship Agreement generated.pdf")
    displayPDF(f1)



def Home(path1, path2, path3):
    # CSS - Points INSTRUCTIONS TO FILL THE AGREEMENT
    st.markdown(""" <style> .font {
        font-size:17px ; font-family: 'Consolas'; color: #0077b6; text-align: justify} 
        </style> """, unsafe_allow_html=True)

    # CSS - Heading INSTRUCTIONS TO FILL THE AGREEMENT
    st.markdown(""" <style> .font1 {
        font-size:26px ; font-family: 'Consolas'; color: #FF0000; text-align: center}
        /style> """, unsafe_allow_html=True)

    st.markdown(""" <style> .font2 {
        font-size:18px ; font-family: 'Consolas'; color: #FF0000; text-align: justify} 
        </style> """, unsafe_allow_html=True)

    # Disclaimer
    st.markdown(""" <style> .font3 {
        font-size:18px ; font-family: 'Consolas'; color: #FF0000; text-align: center} 
        </style> """, unsafe_allow_html=True)

    # Disclaimer Body
    st.markdown(""" <style> .font4 {
        font-size:19px ; font-family: 'Consolas'; color: #FF8C00; text-align: justify} 
        </style> """, unsafe_allow_html=True)

    instruct = st.container()

    with instruct:
        #st.markdown('<p class="font4">Welcome to Law Diktat Agreement builder. We are providing various contracts which the user can fill and use it for their own legal purpose.</p>', unsafe_allow_html=True)
        st.markdown('<p class="font1">INSTRUCTIONS</p>',
                    unsafe_allow_html=True)
        st.markdown('<p class="font">1. After selecting the respective agreement, you will be getting the text boxes to fill in the details for the contract.</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="font">2. You must read the agreement and fill it in accordingly.</p>', unsafe_allow_html=True)
        st.markdown('<p class="font">3. Details entered in the text boxes will be populated in the agreement preview right after pressing the "Enter" or "Tab" key.</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="font">4. Entered text will appear in red-colored texts in the preview box on your right.</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="font">5. The input fields are Case-Sensetive.</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="font">6. You may change the details if required.</p>', unsafe_allow_html=True)
        st.markdown('<p class="font">7. After filling all the fields, you will be able to download the agreement in pdf format.</p>', unsafe_allow_html=True)

        st.sidebar.markdown(
            '<p class="font3"><b><u>DISCLAIMER</u></b></p>', unsafe_allow_html=True)
        #st.sidebar.markdown('<p class="font2"><b>The data you input will be processed by us for our internal usage. We treat your personal information as confidential and will handle it with the utmost care in accordance with the data protection legislation.</b></p>', unsafe_allow_html=True)
        st.sidebar.markdown('<p class="font2"><b>THE DATA YOU INPUT WILL BE PROCESSED BY US FOR OUR INTERNAL USAGE. WE TREAT YOUR PERSONAL INFORMATION AS CONFIDENTIAL AND WILL HANDLE IT WITH THE UTMOST CASE IN ACCORDANCE WITH THE DATA PROTECTION LEGISLATION.</b></p>', unsafe_allow_html=True)


agreements = ['Home', 'General Service Agreement', 'Nullcon Goa Sponsorship Agreement']
agreement_type = st.sidebar.selectbox(
    "SELECT THE AGREEMENT", agreements)
agreement = agreement_type.replace(" ", "")

# pythoncom.CoInitialize()
eval(agreement + "(path1, path2, path3)")
# pythoncom.CoInitialize()
